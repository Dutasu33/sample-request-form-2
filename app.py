# Streamlit 통합 앱 + PDF 출력 + 이메일 발송 기능 포함
import streamlit as st
import pandas as pd
import tempfile
import os
from docx import Document
from docx.shared import Inches
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from datetime import datetime

st.set_page_config(page_title="바디케어 의뢰 + 유사 처방 추천", page_icon="🧴")
st.title("🧴 CNF 제품 개발 의뢰서 (최초 입력용)")

@st.cache_data
def load_data():
    return pd.read_excel("바디워시처방DB20250331.xlsx")

prescription_df = load_data()

st.header("📋 고객 개발 의뢰 정보 입력")
col1, col2 = st.columns(2)
with col1:
    company = st.text_input("회사명")
    manager = st.text_input("담당자명")
    contact = st.text_input("연락처 / 이메일")
    product_name = st.text_input("제품명(가칭)")
    product_type = st.selectbox("제품 유형", ["바디워시", "바디스크럽"])
    volume = st.text_input("용량 / 용기")

with col2:
    launch_date = st.date_input("출시 희망일")
    moq = st.text_input("초도 수량 (MOQ)")
    country = st.text_input("판매 국가")
    scent = st.text_input("희망 향 (예: 시트러스, 머스크 등)")
    texture = st.selectbox("희망 제형", ["오일", "젤", "클레이"])
    is_vegan = st.radio("비건 여부", ["Y", "N"])

if st.button("💾 최초 개발 의뢰서 저장"):
    now = datetime.now().strftime('%Y%m%d_%H%M%S')
    save_path = f"cnf_최초개발의뢰서_{now}.txt"
    with open(save_path, "w", encoding="utf-8") as f:
        f.write(f"회사명: {company}\n담당자명: {manager}\n연락처: {contact}\n")
        f.write(f"제품명: {product_name}\n제품유형: {product_type}\n용량/용기: {volume}\n")
        f.write(f"출시희망일: {launch_date}\n초도수량: {moq}\n판매국가: {country}\n")
        f.write(f"향: {scent}\n제형: {texture}\n비건여부: {is_vegan}\n")
    with open(save_path, "rb") as f:
        st.download_button("📥 저장된 의뢰서 다운로드", f, file_name=save_path)

st.subheader("✍️ 사용감 관련 VOC (자유 서술)")
voc = st.text_area("예: 끈적이지 않고 촉촉하며 향이 오래가는 느낌을 원해요")

st.header("🤖 AI 유사 처방 추천")

if voc:
    corpus = prescription_df['사용감설명'].fillna('').tolist() + [voc]
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(corpus)
    voc_vector = tfidf_matrix[-1]
    prescription_vectors = tfidf_matrix[:-1]

    scores = cosine_similarity(voc_vector, prescription_vectors).flatten()
    prescription_df['유사도'] = scores
    top_df = prescription_df.sort_values(by='유사도', ascending=False).head(3)

    st.success("✅ 추천된 유사 처방 TOP 3")
    st.dataframe(top_df[['처방ID', '제형', '기능성', '주요성분', '향', '보습력', '촉촉함', '제품포지셔닝', '유사도']])

    def generate_docx(data, selected):
        doc = Document()
        doc.add_heading('CNF 확정 제품 개발 의뢰서', 0)

        doc.add_heading('■ 고객사 정보', level=1)
        doc.add_paragraph(f"회사명: {data['회사명']}")
        doc.add_paragraph(f"담당자명: {data['담당자명']}")
        doc.add_paragraph(f"연락처: {data['연락처']}")

        doc.add_heading('■ 제품 기본 정보', level=1)
        doc.add_paragraph(f"제품명(가칭): {data['제품명']}")
        doc.add_paragraph(f"제품유형: {data['제품유형']}")
        doc.add_paragraph(f"용량/용기: {data['용량']}")
        doc.add_paragraph(f"출시희망일: {data['출시일']}")
        doc.add_paragraph(f"초도수량: {data['MOQ']}")
        doc.add_paragraph(f"판매국가: {data['국가']}")

        doc.add_heading('■ 추천 처방 정보', level=1)
        doc.add_paragraph(f"처방랩넘버: {selected['처방ID']}")
        doc.add_paragraph(f"기능성: {selected['기능성']}")
        doc.add_paragraph(f"주요성분: {selected['주요성분']}")
        doc.add_paragraph(f"제품포지셔닝: {selected['제품포지셔닝']}")

        filename = f"cnf_확정개발의뢰서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        doc.save(filepath)
        return filepath

    selected_row = top_df.iloc[0]
    data_dict = {
        "회사명": company, "담당자명": manager, "연락처": contact,
        "제품명": product_name, "제품유형": product_type, "용량": volume,
        "출시일": launch_date.strftime('%Y-%m-%d'), "MOQ": moq, "국가": country
    }
    pdf_path = generate_docx(data_dict, selected_row)

    with open(pdf_path, "rb") as f:
        st.download_button("📄 확정 의뢰서 다운로드 (Word)", f, file_name=os.path.basename(pdf_path))

    st.info("📧 이메일 자동 발송은 실제 서버 환경에서 SMTP 연동이 필요합니다.\nDemo에서는 PDF 다운로드까지만 지원됩니다.")
else:
    st.info("VOC를 입력하면 유사 처방이 자동 추천됩니다.")
